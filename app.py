from flask import Flask, request, jsonify
from docx import Document
from docx.shared import Inches
from filelock import FileLock
import os
import datetime
import pickle
import tensorflow as tf
import numpy as np

app = Flask(__name__)
UPLOAD_DIR = "reports"
TEMP_DIR = "temp"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(TEMP_DIR, exist_ok=True)

# ====================== #
# 🔹 Load model + encoder
# ====================== #
model = tf.keras.models.load_model("bug_severity_vector_model.keras")

# Extract TextVectorization layer from model itself
vectorizer = None
for layer in model.layers:
    if isinstance(layer, tf.keras.layers.TextVectorization):
        vectorizer = layer
        break

if vectorizer is None:
    raise RuntimeError(" TextVectorization layer not found in model!")

# Load label encoder
with open("label_encoder.pkl", "rb") as f:
    label_encoder = pickle.load(f)

# ======================================= #
#  1. Severity Prediction API (AI Part)
# ======================================= #
@app.route('/predict-severity', methods=['POST'])
def predict_severity():
    try:
        data = request.get_json()
        desc = data.get("description", "").lower().strip()

        if not desc:
            return jsonify({"error": "No description provided"}), 400

        # Wrap as string tensor, no external vectorization
        input_tensor = tf.convert_to_tensor([[desc]], dtype=tf.string)
        prediction = model(input_tensor)

        predicted_class = prediction.numpy().argmax(axis=1)[0]
        predicted_label = label_encoder.inverse_transform([predicted_class])[0]

        return jsonify({
            "predicted_severity": predicted_label,
            "confidence": float(np.max(prediction))
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ======================================= #
# 2. Word Document Generation
# ======================================= #
@app.route('/generate-bug-doc', methods=['POST'])
def generate_bug_doc():
    try:
        bug_id = request.form.get("bug_id")
        module = request.form.get("module")
        reporter = request.form.get("reporter")
        assignee = request.form.get("assignee")
        status = request.form.get("status")
        severity = request.form.get("severity")
        description = request.form.get("description")
        filename = request.form.get("filename")
        screenshot = request.files.get("screenshot")

        if not filename:
            return jsonify({"error": "Filename is required"}), 400
        if not filename.endswith(".docx"):
            filename += ".docx"

        file_path = os.path.join(UPLOAD_DIR, filename)
        lock_path = file_path + ".lock"

        with FileLock(lock_path):
            if os.path.exists(file_path):
                doc = Document(file_path)
            else:
                doc = Document()
                doc.add_heading("Bug Report Log", level=1)

            # Check if bug already exists
            bug_found = False
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip() == f"Bug ID: {bug_id}":
                    bug_found = True
                    if i + 1 < len(doc.paragraphs):
                        doc.paragraphs[i + 1].text = f"Time Logged    : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                    if i + 2 < len(doc.paragraphs):
                        doc.paragraphs[i + 2].text = f"Module         : {module}"
                    if i + 3 < len(doc.paragraphs):
                        doc.paragraphs[i + 3].text = f"Reported By    : {reporter}"
                    if i + 4 < len(doc.paragraphs):
                        doc.paragraphs[i + 4].text = f"Assigned To    : {assignee}"
                    if i + 5 < len(doc.paragraphs):
                        doc.paragraphs[i + 5].text = f"Status         : {status}"
                    if i + 6 < len(doc.paragraphs):
                        doc.paragraphs[i + 6].text = f"Severity       : {severity}"
                    if i + 7 < len(doc.paragraphs):
                        doc.paragraphs[i + 7].text = f"Description    : {description}"
                    break

            if not bug_found:
                doc.add_heading(f"Bug ID: {bug_id}", level=2)
                doc.add_paragraph(f"Time Logged    : {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_paragraph(f"Module         : {module}")
                doc.add_paragraph(f"Reported By    : {reporter}")
                doc.add_paragraph(f"Assigned To    : {assignee}")
                doc.add_paragraph(f"Status         : {status}")
                doc.add_paragraph(f"Severity       : {severity}")
                doc.add_paragraph(f"Description    : {description}")

                # Save screenshot if available
                if screenshot and screenshot.filename != "":
                    image_path = os.path.join(TEMP_DIR, f"{bug_id}_screenshot.png")
                    screenshot.save(image_path)
                    doc.add_paragraph("Screenshot:")
                    doc.add_picture(image_path, width=Inches(4))
                    os.remove(image_path)

                doc.add_paragraph("---")

            doc.save(file_path)

        return jsonify({
            "message": f"Bug entry for {bug_id} saved in {filename}",
            "severity": severity
        })

    except Exception as e:
        print(" Word File Error:", e)
        return jsonify({"error": str(e)}), 500

# ==================== #
# Start Flask App
# ==================== #
if __name__ == "__main__":
    app.run(port=5000, debug=True)
